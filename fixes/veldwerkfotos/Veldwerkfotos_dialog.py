import os
import re
import platform
import subprocess
from pathlib import Path
from PIL import Image
from qgis.PyQt import uic, QtWidgets
from qgis.PyQt.QtWidgets import QFileDialog, QMessageBox

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# === Laad het .ui-bestand ===
FORM_CLASS, _ = uic.loadUiType(os.path.join(
    os.path.dirname(__file__), "Veldwerkfotos_dialog_base.ui"))

# === Hulpfuncties ===

def zet_paragraaf_afstand_nul(paragraph):
    pf = paragraph.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1

def verwijder_tabel_randen(table):
    tbl = table._element
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "nil")
        tblBorders.append(border_el)
    tbl.tblPr.append(tblBorders)

def sort_boring_key(boring_nummer):
    num_part = int(re.match(r"^\d+", boring_nummer).group())
    letter_part = re.search(r"[A-Za-z]$", boring_nummer)
    letter = letter_part.group().lower() if letter_part else ""
    return (num_part, letter)

def open_bestand(bestandspad):
    """Open bestand cross-platform"""
    try:
        system = platform.system()
        if system == "Windows":
            os.startfile(bestandspad)
        elif system == "Darwin":  # macOS
            subprocess.call(["open", bestandspad])
        else:  # Linux
            subprocess.call(["xdg-open", bestandspad])
    except Exception as e:
        print(f"❌ Kon bestand niet openen: {e}")

# === Plugin-dialog klasse ===

class VeldwerkfotosDialog(QtWidgets.QDialog, FORM_CLASS):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.btn_selecteer_map.clicked.connect(self.selecteer_foto_map)
        self.btn_genereer_word.clicked.connect(self.genereer_word_document)

    def selecteer_foto_map(self):
        map_pad = QFileDialog.getExistingDirectory(
            self,
            "Selecteer de map met veldwerkfotos",
            "",
            QFileDialog.ShowDirsOnly,
        )
        if map_pad:
            self.lineEdit_foto_map.setText(map_pad)

    def genereer_word_document(self):
        foto_map = self.lineEdit_foto_map.text().strip()

        if not foto_map or not os.path.exists(foto_map):
            QMessageBox.warning(self, "Fout", "Selecteer een geldige map met fotos.")
            return

        try:
            doc = Document()
            title = doc.add_heading("Veldwerkfoto's", level=1)
            title_run = title.runs[0]
            title_run.font.name = "Calibri"
            title_run.font.size = Pt(16)
            title_run.font.color.rgb = RGBColor(38, 113, 144)
            zet_paragraaf_afstand_nul(title)
            doc.add_paragraph()

            # Foto's ophalen
            fotos = [f for f in os.listdir(foto_map)
                     if f.lower().endswith((".jpg", ".jpeg", ".png"))]

            # Classificatie
            overzicht_fotos = [f for f in fotos if re.match(r"^\d{8}", f)]
            overige_fotos = [f for f in fotos if f not in overzicht_fotos]

            # Boring fotos herkennen + duplicaten eruit filteren
            boring_pattern = re.compile(r"^([0-9]+[A-Za-z]?)(?:[-_].*)?$")
            boring_dict = {}
            duplicaten = []
            for f in overige_fotos:
                naam = os.path.splitext(f)[0]
                match = boring_pattern.match(naam)
                if match:
                    boring_nummer = match.group(1)
                    if boring_nummer not in boring_dict:
                        boring_dict[boring_nummer] = f
                    else:
                        duplicaten.append(f)
                # geen match => later onbekende_fotos

            boring_fotos_sorted = [
                boring_dict[k] for k in sorted(boring_dict.keys(), key=sort_boring_key)
            ]

            # Onbekende fotos = niet in boring_dict en niet overzicht
            bekende_fotos = set(overzicht_fotos) | set(boring_dict.values()) | set(duplicaten)
            onbekende_fotos = [f for f in fotos if f not in bekende_fotos]

            # Volledige lijst
            fotos_final = overzicht_fotos + boring_fotos_sorted + onbekende_fotos

            # Tabel aanmaken
            aantal_rijen = ((len(fotos_final) + 1) // 2) * 3
            table = doc.add_table(rows=aantal_rijen, cols=2)
            verwijder_tabel_randen(table)

            fotonummer = 1
            row = 0
            for i in range(0, len(fotos_final), 2):
                for col in range(2):
                    if i + col < len(fotos_final):
                        foto_bestand = fotos_final[i + col]
                        foto_pad = os.path.join(foto_map, foto_bestand)
                        naam_zonder_ext = os.path.splitext(foto_bestand)[0]

                        try:
                            with Image.open(foto_pad) as img:
                                orig_width, orig_height = img.size
                                schaal_factor = (
                                    7.5 / (orig_width / 96 * 2.54) if orig_width > 0 else 1
                                )
                                hoogte_cm = max(
                                    1, (orig_height / 96 * 2.54) * schaal_factor
                                )
                        except Exception as e:
                            print(f"❌ Foto '{foto_bestand}' overgeslagen (fout: {e})")
                            continue

                        # Foto
                        cell = table.cell(row, col)
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(foto_pad, width=Cm(7.5), height=Cm(hoogte_cm))
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        zet_paragraaf_afstand_nul(p)

                        # Fotonummer
                        cell = table.cell(row + 1, col)
                        p2 = cell.paragraphs[0]
                        r2 = p2.add_run(f"Fotonummer: {fotonummer}")
                        r2.font.name = "Calibri"
                        r2.font.size = Pt(8)
                        r2.font.color.rgb = RGBColor(38, 113, 144)
                        r2.bold = True
                        p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        zet_paragraaf_afstand_nul(p2)

                        # Omschrijving
                        cell = table.cell(row + 2, col)
                        p3 = cell.paragraphs[0]

                        if foto_bestand in overzicht_fotos:
                            omschrijving = "Omschrijving: Overzicht onderzoeksgebied"
                            kleur = RGBColor(0, 0, 0)
                            bold = False
                            run = p3.add_run(omschrijving)
                            run.font.name = "Calibri"
                            run.font.size = Pt(8)
                            run.font.color.rgb = kleur
                            run.bold = bold

                        elif foto_bestand in boring_dict.values():
                            boring_nummer = re.split(r"[-_]", naam_zonder_ext)[0]
                            omschrijving = f"Omschrijving: Boring {boring_nummer}"
                            kleur = RGBColor(0, 0, 0)
                            bold = False
                            run = p3.add_run(omschrijving)
                            run.font.name = "Calibri"
                            run.font.size = Pt(8)
                            run.font.color.rgb = kleur
                            run.bold = bold

                        else:  # onbekend
                            # "Omschrijving: " zwart en normaal
                            run1 = p3.add_run("Omschrijving: ")
                            run1.font.name = "Calibri"
                            run1.font.size = Pt(8)
                            run1.font.color.rgb = RGBColor(0, 0, 0)
                            run1.bold = False

                            # bestandsnaam rood en vet
                            run2 = p3.add_run(naam_zonder_ext)
                            run2.font.name = "Calibri"
                            run2.font.size = Pt(8)
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                            run2.bold = True

                        p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                        fotonummer += 1
                row += 3

            # Document opslaan
            downloads_dir = Path.home() / "Downloads"
            save_path = downloads_dir / "Bijlage veldwerkfoto's.docx"
            doc.save(str(save_path))

            # Logbestand aanmaken indien nodig
            # Vervang dit stuk in het originele script (onder het logbestand aanmaken):

            log_info = []
            if duplicaten or onbekende_fotos:
                log_path = downloads_dir / "Bijlage veldwerkfoto's_log.txt"
                with open(log_path, "w", encoding="utf-8") as log:
                    if duplicaten:
                        log.write("Duplicaten:\n")
                        log.writelines(f" - {d}\n" for d in duplicaten)
                        aantal = len(duplicaten)
                        tekst = f"{aantal} duplicaat" if aantal == 1 else f"{aantal} duplicaten"
                        log_info.append(tekst)
                    if onbekende_fotos:
                        log.write("\nOnbekende foto's:\n")
                        log.writelines(f" - {o}\n" for o in onbekende_fotos)
                        aantal = len(onbekende_fotos)
                        tekst = f"{aantal} onbekende foto" if aantal == 1 else f"{aantal} onbekende foto's"
                        log_info.append(tekst)

            melding = f"Document opgeslagen als:\n{save_path}"
            if log_info:
                melding += "\n\n" + ", ".join(log_info) + " (zie logbestand)."

            QMessageBox.information(self, "Succes", melding)

            # Open document na OK
            open_bestand(str(save_path))

        except Exception as e:
            QMessageBox.critical(self, "Fout", f"Er is iets misgegaan:\n{e}")
